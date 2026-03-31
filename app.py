import streamlit as st
from docx import Document
from io import BytesIO
import re
import calendar
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import pandas as pd

# --- CORE FUNCTIONS ---
def fill_template(template, data):
    doc = Document(template)
    def replace_text_in_paragraph(paragraph, data):
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(str(key), str(value))
    
    # Standard Paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, data)
    
    # Tables and Nested Cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, data)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def format_suffix_date(d):
    day = d.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return d.strftime(f"{day}{suffix} %B %Y")

def number_to_word_format(n_str):
    words = {
        '1': 'One', '2': 'Two', '3': 'Three', '4': 'Four', '5': 'Five', '6': 'Six', '7': 'Seven', '8': 'Eight', '9': 'Nine', '10': 'Ten',
        '11': 'Eleven', '12': 'Twelve', '13': 'Thirteen', '14': 'Fourteen', '15': 'Fifteen', '16': 'Sixteen', '17': 'Seventeen', '18': 'Eighteen', '19': 'Nineteen', '20': 'Twenty',
        '21': 'Twenty-One', '22': 'Twenty-Two', '23': 'Twenty-Three', '24': 'Twenty-Four', '25': 'Twenty-Five', '26': 'Twenty-Six', '27': 'Twenty-Seven', '28': 'Twenty-Eight', '29': 'Twenty-Nine', '30': 'Thirty',
        '31': 'Thirty-One', '32': 'Thirty-Two', '33': 'Thirty-Three', '34': 'Thirty-Four', '35': 'Thirty-Five', '36': 'Thirty-Six', '37': 'Thirty-Seven', '38': 'Thirty-Eight', '39': 'Thirty-Nine', '40': 'Forty',
        '41': 'Forty-One', '42': 'Forty-Two', '43': 'Forty-Three', '44': 'Forty-Four', '45': 'Forty-Five', '46': 'Forty-Six', '47': 'Forty-Seven', '48': 'Forty-Eight', '49': 'Forty-Nine', '50': 'Fifty',
        '51': 'Fifty-One', '52': 'Fifty-Two', '53': 'Fifty-Three', '54': 'Fifty-Four', '55': 'Fifty-Five', '56': 'Fifty-Six', '57': 'Fifty-Seven', '58': 'Fifty-Eight', '59': 'Fifty-Nine', '60': 'Sixty'
    }
    word = words.get(str(n_str), str(n_str)) 
    return f"{word} ({n_str})"

# --- SESSION STATE ---
if 'show_warnings' not in st.session_state:
    st.session_state.show_warnings = False
if 'generated' not in st.session_state:
    st.session_state.generated = False

def clear_form():
    st.session_state.generated = False
    st.session_state.show_warnings = False
    for key in list(st.session_state.keys()):
        if key.startswith("in_"):
            if "date" in key:
                st.session_state[key] = datetime.today().date()
            else:
                st.session_state[key] = ""

st.set_page_config(page_title="LoE Generator", layout="centered")

st.markdown("<h1 style='font-size: 36px; font-weight: bold;'>Step 1: Upload Templates</h1>", unsafe_allow_html=True)
col_t1, col_t2 = st.columns(2)
with col_t1:
    perm_template = st.file_uploader("Permanent Template", type="docx", key="ft_perm")
with col_t2:
    fix_template = st.file_uploader("Fixed Term Template", type="docx", key="ftc_perm")

if perm_template or fix_template:
    st.divider()
    st.subheader("Step 2: Category & Details")
    emp_type = st.radio("Generate letter for:", ["Permanent Employment", "Fixed Term"], horizontal=True)
    selected_template = perm_template if emp_type == "Permanent Employment" else fix_template
    
    if not selected_template:
        st.error(f"❌ Please upload the {emp_type} template first.")
    else:
        v = {} # Master Validation Tracker

        col_main_l, col_main_r = st.columns(2)
        with col_main_l:
            # Full Name
            name = st.text_input("Full Name*", key="in_name")
            v['name'] = bool(name and re.fullmatch(r"^[a-zA-Z\s]+$", name))
            if name and not v['name']: st.error("⚠️ Full Name: Letters and spaces only.")
            elif not name and st.session_state.show_warnings: st.warning("⚠️ Full Name is required.")

            # Applicant Address (Permanent Only)
            emp_address = ""
            v['address'] = True
            if emp_type == "Permanent Employment":
                emp_address = st.text_area("Applicant Address*", key="in_address")
                v['address'] = bool(emp_address.strip())
                if not v['address'] and st.session_state.show_warnings: st.warning("⚠️ Applicant Address is required.")

            # IC Number
            ic_number = st.text_input("IC Number (12 Digits)*", max_chars=12, key="in_ic")
            v['ic'] = bool(ic_number.isdigit() and len(ic_number) == 12)
            if ic_number and not v['ic']: st.error("⚠️ IC Number: Exactly 12 digits required.")
            elif not ic_number and st.session_state.show_warnings: st.warning("⚠️ IC Number is required.")

            # Job Title
            job_title = st.text_input("Job Title*", key="in_job")
            v['job'] = bool(job_title)
            if not job_title and st.session_state.show_warnings: st.warning("⚠️ Job Title is required.")

        with col_main_r:
            # Project Name (Fixed Term Only)
            project = ""
            v['project'] = True
            if emp_type == "Fixed Term":
                project = st.text_input("Project Name*", key="in_project")
                v['project'] = bool(project and re.fullmatch(r"^[a-zA-Z\s]+$", project))
                if project and not v['project']: st.error("⚠️ Project Name: Letters and spaces only.")
                elif not project and st.session_state.show_warnings: st.warning("⚠️ Project Name is required.")

            # Employee Grade
            emp_grade = st.text_input("Employee Grade*", key="in_grade")
            v['grade'] = bool(emp_grade)
            if not emp_grade and st.session_state.show_warnings: st.warning("⚠️ Employee Grade is required.")

            # Salary
            s1, s2 = st.columns([0.5, 3.5])
            s1.markdown("<br>RM", unsafe_allow_html=True)
            salary_raw = s2.text_input("Monthly Salary*", key="in_salary")
            salary_clean = salary_raw.replace(",", "")
            v['salary'] = bool(re.match(r"^\d+(\.\d{1,2})?$", salary_clean)) if salary_clean else False
            if salary_raw and not v['salary']: st.error("⚠️ Salary: Use format 1,000.00")
            elif not salary_raw and st.session_state.show_warnings: st.warning("⚠️ Monthly Salary is required.")
            
            # CEO Name
            ceo_name = st.text_input("CEO Name*", value="SAMANTHA TAN", key="in_ceo")
            v['ceo'] = bool(ceo_name and re.fullmatch(r"^[a-zA-Z\s]+$", ceo_name))
            if ceo_name and not v['ceo']: st.error("⚠️ CEO Name: Letters and spaces only.")
            elif not ceo_name and st.session_state.show_warnings: st.warning("⚠️ CEO Name is required.")
            
            start_date = st.date_input("Start Date*", key="in_start_date")

        st.markdown("#### Contractual Terms")
        cl, cr = st.columns(2)
        with cl:
            # Probation
            m1, m2 = cl.columns([3, 2])
            probation = m1.text_input("Probation*", key="in_prob")
            m2.markdown("<br>Months", unsafe_allow_html=True)
            v['prob'] = bool(probation.isdigit())
            if probation and not v['prob']: st.error("⚠️ Probation: Digits only.")
            elif not probation and st.session_state.show_warnings: st.warning("⚠️ Probation period required.")

            # Notice - Probation
            np1, np2 = cl.columns([3, 2])
            n_prob = np1.text_input("Notice - Probation*", key="in_nprob")
            np2.markdown("<br>Months", unsafe_allow_html=True)
            v['nprob'] = bool(n_prob.isdigit())
            if n_prob and not v['nprob']: st.error("⚠️ Notice Probation: Digits only.")
            elif not n_prob and st.session_state.show_warnings: st.warning("⚠️ Notice (Probation) required.")

            # Notice - Confirmed
            nc1, nc2 = cl.columns([3, 2])
            n_conf = nc1.text_input("Notice - Confirmed*", key="in_nconf")
            nc2.markdown("<br>Months", unsafe_allow_html=True)
            v['nconf'] = bool(n_conf.isdigit())
            if n_conf and not v['nconf']: st.error("⚠️ Notice Confirmed: Digits only.")
            elif not n_conf and st.session_state.show_warnings: st.warning("⚠️ Notice (Confirmed) required.")

        with cr:
            # Annual Leave
            al1, al2 = cr.columns([3, 2])
            al_val = al1.text_input("Annual Leave*", key="in_al")
            al2.markdown("<br>Days", unsafe_allow_html=True)
            v['al'] = bool(al_val.isdigit())
            if al_val and not v['al']: st.error("⚠️ Annual Leave: Digits only.")
            elif not al_val and st.session_state.show_warnings: st.warning("⚠️ Annual Leave required.")

            # Outpatient Limit
            o1, o2 = cr.columns([0.5, 3.5])
            o1.markdown("<br>RM", unsafe_allow_html=True)
            outpatient = o2.text_input("Outpatient Limit*", value=0, key="in_out")
            v['out'] = bool(re.match(r"^\d+(\.\d{1,2})?$", str(outpatient).replace(",","")))
            if outpatient and not v['out']: st.error("⚠️ Outpatient: Digits/Decimals only.")
            
            # Travel & KPI (Fixed Term Only)
            travel, kpi = "0", "0"
            v['travel'], v['kpi'] = True, True
            if emp_type == "Fixed Term":
                tr1, tr2 = cr.columns([0.5, 3.5])
                tr1.markdown("<br>RM", unsafe_allow_html=True)
                travel = tr2.text_input("Traveling Allowance*", value=0, key="in_travel")
                v['travel'] = bool(re.match(r"^\d+(\.\d{1,2})?$", str(travel).replace(",","")))
                if travel and not v['travel']: st.error("⚠️ Travel Allowance: Digits only.")
                elif not travel and st.session_state.show_warnings: st.warning("⚠️ Travel Allowance required.")

                kp1, kp2 = cr.columns([0.5, 3.5])
                kp1.markdown("<br>RM", unsafe_allow_html=True)
                kpi = kp2.text_input("Max KPI Payout*", value=0, key="in_kpi")
                v['kpi'] = bool(re.match(r"^\d+(\.\d{1,2})?$", str(kpi).replace(",","")))
                if kpi and not v['kpi']: st.error("⚠️ KPI: Digits only.")
                elif not kpi and st.session_state.show_warnings: st.warning("⚠️ Max KPI Payout required.")

        v['fix'] = True
        if emp_type == "Fixed Term":
            st.markdown("#### Transitional & Term Details")
            t_c1, t_c2 = st.columns([3, 2])
            c_term = t_c1.text_input("Contract Term*", key="in_fix_term")
            t_c2.markdown("<br>Months", unsafe_allow_html=True)
            v['term'] = bool(c_term.isdigit())
            if c_term and not v['term']: st.error("⚠️ Contract Term: Digits only.")
            elif not c_term and st.session_state.show_warnings: st.warning("⚠️ Contract Term required.")
            
            months_list = list(calendar.month_name)[1:]
            years_list = [datetime.today().year + i for i in range(-1, 3)]
            m1, y1 = st.columns(2)
            t_sm_name = m1.selectbox("Trans. Start Month", months_list, key="in_fix_tsm")
            t_sy = y1.selectbox("Start Year", years_list, index=1, key="in_fix_tyear")
            m2, y2 = st.columns(2)
            t_em_name = m2.selectbox("Trans. End Month", months_list, key="in_fix_tem")
            t_ey = y2.selectbox("End Year", years_list, index=1, key="in_fix_teyear")
            
            f_ts = datetime(t_sy, months_list.index(t_sm_name) + 1, 1)
            f_te = datetime(t_ey, months_list.index(t_em_name) + 1, calendar.monthrange(t_ey, months_list.index(t_em_name) + 1)[1])
            v['dates'] = f_te >= f_ts
            if not v['dates']: st.error("⚠️ End Month must be after Start Month.")
            
            f_as = f_te + timedelta(days=1)
            prev_expiry_calc = f_ts - timedelta(days=1)
            trans_month_label = f_ts.strftime("%B %Y")
            v['fix'] = v['term'] and v['dates']

        # MASTER READINESS CHECK
        is_ready = all(v.values())

        st.divider()
        if is_ready:
            if st.button(f"🚀 Generate {emp_type} Letter", type="primary"):
                st.session_state.generated = True
        else:
            if st.button(f"🚀 Check for Errors", type="secondary"):
                st.session_state.show_warnings = True
                st.rerun()
            st.info("💡 Please fix all errors or fill required fields (*) to enable the Generate button.")

        if st.session_state.generated and is_ready:
            with st.expander("✅ Review Data & Download", expanded=True):
                # Calculations
                prob_months = int(probation) if probation.isdigit() else 0
                days_to_add = prob_months * 30
                p_end_raw = start_date + timedelta(days=days_to_add - 1)
                
                def fmt_rm(v_val):
                    clean = str(v_val).replace(",","")
                    if not clean or float(clean) == 0: return "RM TBA"
                    return f"RM {float(clean):,.2f}"

                data_map = {
                    "{{TODAY}}": format_suffix_date(datetime.today()),
                    "{{NAME}}": name.upper(), "{{IC_NUMBER}}": ic_number,
                    "{{JOB_TITLE}}": job_title, "{{GRADE}}": emp_grade.upper(), 
                    "{{START_DATE}}": format_suffix_date(start_date),
                    "{{SALARY}}": fmt_rm(salary_clean), "{{CEO_NAME}}": ceo_name.upper(),
                    "{{PROBATION}}": probation, "{{PROBATION_DAYS}}": str(days_to_add),
                    "{{PROBATION_END}}": format_suffix_date(p_end_raw),
                    "{{NOTICE_PROB}}": n_prob, "{{NOTICE_CONF}}": n_conf, 
                    "{{AL}}": al_val, "{{OUTPATIENT}}": fmt_rm(outpatient)
                }

                if emp_type == "Permanent Employment":
                    data_map["{{EMP_ADDRESS}}"] = emp_address.strip()
                else:
                    data_map.update({
                        "{{PROJECT}}": project.upper(), "{{TRAVEL}}": fmt_rm(travel),
                        "{{KPI}}": fmt_rm(kpi), "{{TERM}}": number_to_word_format(c_term),
                        "{{EXPIRY_DATE}}": format_suffix_date(start_date + relativedelta(months=int(c_term))),
                        "{{TRANS_START}}": format_suffix_date(f_ts), "{{TRANS_END}}": format_suffix_date(f_te),
                        "{{AGREE_START}}": format_suffix_date(f_as), "{{PREV_EXPIRY}}": format_suffix_date(prev_expiry_calc),
                        "{{TRANS_MONTH}}": trans_month_label
                    })

                type_code = "PERMANENT" if emp_type[0]=='P' else "FIXED-TERM"
                fname = f"LOE_{type_code}_{name.upper().replace(' ','_')}_{start_date.strftime('%Y_%m_%d')}.docx"
                st.table(pd.DataFrame(list(data_map.items()), columns=["Tag", "Value"]))
                st.download_button(f"📥 Download {fname}", data=fill_template(selected_template, data_map), file_name=fname)
                st.button("🗑️ Reset Form", on_click=clear_form)
